"""Convert a subset of Markdown to a .docx file using python-docx.

Supported constructs:
  - # / ## / ### headings  → Heading 1 / 2 / 3
  - - / * bullet lists     → List Bullet
  - GitHub-style pipe tables (| a | b | with |---|) → real docx Table
  - **bold** inline        → Bold run (best-effort, no nested formatting)
  - Everything else        → Normal paragraph

Unknown or malformed Markdown degrades silently to plain paragraphs —
this function never raises on odd input.
"""

from __future__ import annotations

import base64
import binascii
import io
import logging
import re
from typing import TYPE_CHECKING

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches

if TYPE_CHECKING:
    from docx.document import Document as DocxDocument

logger = logging.getLogger(__name__)

# PNG files start with this 8-byte signature.
_PNG_SIGNATURE = b"\x89PNG\r\n\x1a\n"
_DATA_URL_PREFIX_RE = re.compile(r"^data:image/[^;]+;base64,", re.IGNORECASE)

# ── regex helpers ──────────────────────────────────────────────────────────────

_HEADING_RE = re.compile(r"^(#{1,3})\s+(.+)$")
_BULLET_RE = re.compile(r"^[-*]\s+(.+)$")
_TABLE_SEP_RE = re.compile(r"^\|[-| :]+\|$")
_TABLE_ROW_RE = re.compile(r"^\|(.+)\|$")
_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
# Markdown image: ![alt](url). The agent emits ![Gráfico …](link_del_grafico)
# as a placeholder for where the chart should sit. We never render the literal
# syntax — we either embed the real chart there or drop the line entirely.
_IMAGE_RE = re.compile(r"^!\[[^\]]*\]\([^)]*\)$")


def _add_runs_with_bold(para: object, text: str) -> None:
    """Split *text* on **…** markers and add runs, bolding the marked parts."""
    parts = _BOLD_RE.split(text)
    # split gives alternating [plain, bold, plain, bold, …]
    for i, part in enumerate(parts):
        if not part:
            continue
        run = para.add_run(part)  # type: ignore[union-attr]
        run.bold = bool(i % 2 == 1)


def _parse_pipe_row(line: str) -> list[str]:
    """Extract cell text from a pipe-table row, stripping surrounding whitespace."""
    inner = line.strip().strip("|")
    return [cell.strip() for cell in inner.split("|")]


def _shade_cell(cell: object, fill_hex: str) -> None:
    """Apply a solid shading to a table cell (hex without #, e.g. '4F81BD')."""
    tc = cell._tc  # type: ignore[union-attr]
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _build_table(doc: DocxDocument, rows: list[list[str]]) -> None:
    """Add a styled table to *doc* from a list of cell-text rows.

    Row 0 is the header (shaded blue, bold white text).
    Remaining rows alternate light-grey / white.
    """
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    # Pad all rows to the same width
    padded = [r + [""] * (n_cols - len(r)) for r in rows]

    table = doc.add_table(rows=len(padded), cols=n_cols)
    table.style = "Table Grid"

    for row_idx, row_data in enumerate(padded):
        tr = table.rows[row_idx]
        is_header = row_idx == 0
        for col_idx, cell_text in enumerate(row_data):
            cell = tr.cells[col_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            if is_header:
                run.bold = True
                run.font.color.rgb = None  # keep auto (will show on dark bg)
                _shade_cell(cell, "4F81BD")  # medium blue header
            elif row_idx % 2 == 0:
                _shade_cell(cell, "DCE6F1")  # light blue stripe
            # else: white (default)


def _slug(titulo: str) -> str:
    """Turn a title into a safe ASCII filename stem (max 60 chars)."""
    slug = re.sub(r"[^a-zA-Z0-9\-_\s]", "", titulo.lower())
    slug = re.sub(r"\s+", "-", slug.strip())
    return slug[:60] or "documento"


def _decode_png(chart_image_base64: str) -> bytes | None:
    """Decode a base64 PNG (data URL or bare) to bytes, or None if invalid.

    Returns None — never raises — when the input is malformed or not a PNG.
    """
    raw = _DATA_URL_PREFIX_RE.sub("", chart_image_base64.strip(), count=1)
    try:
        png_bytes = base64.b64decode(raw, validate=True)
    except (binascii.Error, ValueError):
        logger.warning("chart_image_base64 is not valid base64 — skipping image")
        return None
    if not png_bytes.startswith(_PNG_SIGNATURE):
        logger.warning("chart_image_base64 did not decode to a PNG — skipping image")
        return None
    return png_bytes


def _embed_png(doc: DocxDocument, png_bytes: bytes, *, heading: str | None = None) -> bool:
    """Embed a decoded PNG into *doc* (optionally under a heading).

    Returns True on success, False if python-docx rejected the image — never
    raises, so a bad image can't 500 the whole document.
    """
    try:
        if heading:
            doc.add_heading(heading, level=2)
        doc.add_picture(io.BytesIO(png_bytes), width=Inches(6))
        return True
    except Exception:  # never 500 the whole document over an image
        logger.warning("failed to embed chart image into docx — skipping", exc_info=True)
        return False


# ── main entry point ───────────────────────────────────────────────────────────


def render(
    titulo: str,
    contenido_markdown: str,
    *,
    chart_image_base64: str | None = None,
    include_tables: bool = True,
) -> bytes:
    """Convert *titulo* + *contenido_markdown* to .docx bytes.

    The chart (when *chart_image_base64* is a valid PNG) is embedded at the
    first ``![…](…)`` image placeholder the agent wrote — that's where it
    belongs in the narrative. If the document has no placeholder, the chart is
    appended under a 'Gráfico' heading instead. When no chart is supplied, image
    placeholders are dropped (never rendered as literal Markdown).

    When *include_tables* is False, pipe tables are skipped entirely. Bad
    markdown or a bad image degrade silently — this function never raises.
    """
    doc: DocxDocument = Document()

    # Document title as Heading 1 at the top
    doc.add_heading(titulo, level=1)

    chart_png = _decode_png(chart_image_base64) if chart_image_base64 else None
    chart_embedded = False

    lines = contenido_markdown.splitlines()
    i = 0
    pending_table: list[list[str]] = []

    def _flush_table() -> None:
        nonlocal pending_table
        if pending_table:
            _build_table(doc, pending_table)
            pending_table = []

    while i < len(lines):
        # rstrip so trailing whitespace (common in LLM-generated tables, where
        # rows end with "| ") doesn't defeat the anchored `…\|$` table regexes.
        line = lines[i].rstrip()

        # ── heading ──────────────────────────────────────────────────────────
        m_h = _HEADING_RE.match(line)
        if m_h:
            _flush_table()
            level = min(len(m_h.group(1)), 3)
            doc.add_heading(m_h.group(2).strip(), level=level)
            i += 1
            continue

        # ── image placeholder ![alt](url) ─────────────────────────────────────
        if _IMAGE_RE.match(line):
            _flush_table()
            if chart_png is not None and not chart_embedded:
                chart_embedded = _embed_png(doc, chart_png)
            # else: drop the placeholder — never render literal Markdown syntax
            i += 1
            continue

        # ── pipe table row ───────────────────────────────────────────────────
        if _TABLE_ROW_RE.match(line):
            if not include_tables:
                i += 1  # caller asked to hide tables — skip every table line
                continue
            # separator row (|---|) — skip, it's already handled by the header
            if _TABLE_SEP_RE.match(line):
                i += 1
                continue
            pending_table.append(_parse_pipe_row(line))
            i += 1
            continue

        # If we were building a table and hit a non-table line, flush it first
        _flush_table()

        # ── bullet ───────────────────────────────────────────────────────────
        m_b = _BULLET_RE.match(line)
        if m_b:
            para = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(para, m_b.group(1))
            i += 1
            continue

        # ── blank line ───────────────────────────────────────────────────────
        if not line.strip():
            i += 1
            continue

        # ── default: plain paragraph ─────────────────────────────────────────
        para = doc.add_paragraph()
        _add_runs_with_bold(para, line)
        i += 1

    _flush_table()

    # Chart supplied but the document had no placeholder to host it → append.
    if chart_png is not None and not chart_embedded:
        _embed_png(doc, chart_png, heading="Gráfico")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def filename_for(titulo: str) -> str:
    """Return the attachment filename that the endpoint should advertise."""
    return f"{_slug(titulo)}.docx"
