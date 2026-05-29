"""Append the AI/NLP analysis sections to a claim's Word report.

Kept separate from `generate_claim_report_docx` so the main builder stays under
the 250-LOC convention and this module owns a single concern: rendering the
explainability surfaces (ML factors, anomaly, similar narratives, NLP narrative
read, multi-agent panel) that the analyst sees on the claim detail page.

Framing rule (§2.10): every section is *análisis* / *alerta*, never an
accusation. The panel + NLP outputs are advisory and never override the score.
"""

from __future__ import annotations

from docx.document import Document

from app.schemas.claim import ClaimDetail
from app.schemas.risk import Tier

_TIER_LABELS: dict[Tier, str] = {
    Tier.verde: "Bajo",
    Tier.amarillo: "Medio",
    Tier.rojo: "Alto",
}


def append_analysis_sections(doc: Document, claim: ClaimDetail) -> None:
    """Append every available analysis section, in detail-page order."""
    _append_ml_factors(doc, claim)
    _append_anomaly(doc, claim)
    _append_similar(doc, claim)
    _append_narrative_analysis(doc, claim)
    _append_panel_analysis(doc, claim)


def _append_ml_factors(doc: Document, claim: ClaimDetail) -> None:
    if not claim.ml_factors:
        return
    doc.add_heading("Factores del modelo (SHAP)", level=1)
    for factor in claim.ml_factors:
        sentido = "aumenta el riesgo" if factor.direction == "up" else "reduce el riesgo"
        doc.add_paragraph(
            f"{factor.feature}: {sentido} (SHAP {factor.shap_value:.3f})",
            style="List Bullet",
        )
    doc.add_paragraph("")


def _append_anomaly(doc: Document, claim: ClaimDetail) -> None:
    if claim.anomaly_score is None:
        return
    doc.add_heading("Indicador de anomalía", level=1)
    p = doc.add_paragraph()
    p.add_run("Anomaly score: ").bold = True
    p.add_run(f"{claim.anomaly_score:.3f} (valores más bajos = más atípico)")
    if claim.nearest_normal_claim_id:
        doc.add_paragraph(
            f"Caso normal más cercano para contraste: {claim.nearest_normal_claim_id}"
        )
    doc.add_paragraph("")


def _append_similar(doc: Document, claim: ClaimDetail) -> None:
    if not claim.similar:
        return
    doc.add_heading("Narrativas similares", level=1)
    for s in claim.similar:
        doc.add_paragraph(
            f"{s.claim_id} — similitud {s.similarity * 100:.1f}%: {s.snippet}",
            style="List Bullet",
        )
    doc.add_paragraph("")


def _append_narrative_analysis(doc: Document, claim: ClaimDetail) -> None:
    na = claim.narrative_analysis
    if na is None:
        return
    doc.add_heading("Análisis NLP del relato", level=1)
    if na.resumen_narrativa:
        doc.add_paragraph(na.resumen_narrativa)

    if na.narrativa_ilogica or na.incoherencias:
        p = doc.add_paragraph()
        p.add_run("Posibles incoherencias detectadas:").bold = True
        for inc in na.incoherencias:
            doc.add_paragraph(inc, style="List Bullet")
    else:
        doc.add_paragraph("El relato no presenta incoherencias internas evidentes.")

    entidades = na.entidades
    grupos = [
        ("Personas", entidades.personas),
        ("Lugares", entidades.lugares),
        ("Fechas", entidades.fechas),
        ("Vehículos", entidades.vehiculos),
        ("Terceros", entidades.terceros),
        ("Montos", entidades.montos),
    ]
    presentes = [(etiqueta, valores) for etiqueta, valores in grupos if valores]
    if presentes:
        p = doc.add_paragraph()
        p.add_run("Entidades extraídas:").bold = True
        for etiqueta, valores in presentes:
            doc.add_paragraph(f"{etiqueta}: {', '.join(valores)}", style="List Bullet")
    doc.add_paragraph("")


def _append_panel_analysis(doc: Document, claim: ClaimDetail) -> None:
    panel = claim.panel_analysis
    if panel is None:
        return
    doc.add_heading("Panel multi-agente", level=1)
    intro = doc.add_paragraph()
    intro.add_run(
        "Análisis colaborativo de especialistas. Es advisory — no modifica el "
        "score ni constituye una decisión."
    ).italic = True

    consensus = panel.consensus
    if consensus is not None:
        nivel = _TIER_LABELS.get(consensus.nivel_final, consensus.nivel_final.value)
        p = doc.add_paragraph()
        p.add_run("Consenso: ").bold = True
        p.add_run(
            f"nivel {nivel} · acuerdo {consensus.nivel_de_acuerdo * 100:.0f}%"
        )
        if consensus.resumen:
            doc.add_paragraph(consensus.resumen)
        if consensus.accion_recomendada:
            p = doc.add_paragraph()
            p.add_run("Acción recomendada: ").bold = True
            p.add_run(consensus.accion_recomendada)
        if consensus.posible_falso_positivo:
            doc.add_paragraph("El panel señala un posible falso positivo.")
        for conflicto in consensus.puntos_de_conflicto:
            doc.add_paragraph(f"Conflicto: {conflicto}", style="List Bullet")

    if panel.moderator_text:
        p = doc.add_paragraph()
        p.add_run("Moderador: ").bold = True
        p.add_run(panel.moderator_text)

    for lane in panel.lanes:
        if lane.failed:
            continue
        doc.add_heading(f"{lane.display_name} · {lane.lens}", level=2)
        verdict = lane.verdict
        if verdict is not None:
            nivel = _TIER_LABELS.get(verdict.nivel, verdict.nivel.value)
            p = doc.add_paragraph()
            p.add_run(f"Nivel {nivel}: ").bold = True
            p.add_run(verdict.dictamen)
            for punto in verdict.puntos_clave:
                doc.add_paragraph(punto, style="List Bullet")
    doc.add_paragraph("")
