"""Generate a Word (.docx) report for a single claim.

Sections
--------
1. Título — "Reporte de caso {id}"
2. Datos del caso — asegurado, póliza, cobertura, ciudad, montos, fechas
3. Score + nivel — traffic-light wording, never "fraude" without "posible"
4. Reglas activadas — table: código / puntos / detalle
4b. Análisis explicativo — ML factors / anomalía / narrativas similares / NLP /
    panel multi-agente (delegated to `append_analysis_sections`)
5. Resumen — resumen_editado when present, else data-driven fallback
6. Recomendación de revisión

Design notes
-----------
- docx is CPU-bound; wrapped in asyncio.to_thread to avoid blocking the event loop.
- Returns raw bytes; the caller wraps into a StreamingResponse / Response.
- < 250 LOC per convention.
"""

from __future__ import annotations

import asyncio
from io import BytesIO

from docx import Document  # type: ignore[import-untyped]
from docx.shared import Pt, RGBColor  # type: ignore[import-untyped]

from app.schemas.claim import ClaimDetail
from app.schemas.risk import Tier
from app.use_cases._claim_report_analysis import append_analysis_sections

# Traffic-light wording — never "fraude" without "posible" (CLAUDE.md §2 / §17)
_NIVEL_LABELS: dict[Tier, str] = {
    Tier.verde: "Verde — flujo normal (0-40 puntos)",
    Tier.amarillo: "Amarillo — requiere revisión de documentos (41-75 puntos)",
    Tier.rojo: "Rojo — requiere revisión en campo por posible fraude (76-100 puntos)",
}

_NIVEL_COLORS: dict[Tier, tuple[int, int, int]] = {
    Tier.verde: (0, 128, 0),
    Tier.amarillo: (204, 153, 0),
    Tier.rojo: (192, 0, 0),
}

_RECOMENDACION: dict[Tier, str] = {
    Tier.verde: (
        "El caso no presenta alertas críticas. Se recomienda seguir el flujo "
        "normal de procesamiento."
    ),
    Tier.amarillo: (
        "El caso presenta señales de alerta. Se recomienda escalar a la Unidad "
        "Antifraude para revisión de documentos antes de proceder."
    ),
    Tier.rojo: (
        "El caso presenta múltiples alertas de alto riesgo de posible fraude. "
        "Se recomienda escalar con carácter urgente a la Unidad Antifraude "
        "para revisión en campo. El analista debe retener el pago hasta que se "
        "complete la investigación."
    ),
}


def _build_docx(claim: ClaimDetail) -> bytes:
    """Synchronous builder — runs inside asyncio.to_thread."""
    doc = Document()

    # -----------------------------------------------------------------------
    # 1. Título
    # -----------------------------------------------------------------------
    title = doc.add_heading(f"Reporte de caso {claim.id}", level=0)
    title.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "Este documento es confidencial y está destinado exclusivamente al "
        "personal autorizado de la Unidad de Análisis de la Aseguradora. "
        "Las alertas contenidas son indicativas — no constituyen una acusación."
    ).italic = True

    doc.add_paragraph("")

    # -----------------------------------------------------------------------
    # 2. Datos del caso
    # -----------------------------------------------------------------------
    doc.add_heading("Datos del caso", level=1)

    def _row(label: str, value: str) -> None:
        p = doc.add_paragraph()
        run_label = p.add_run(f"{label}: ")
        run_label.bold = True
        p.add_run(value)

    _row("Asegurado", claim.asegurado)
    _row("ID asegurado", claim.asegurado_id)
    _row("Póliza", claim.poliza)
    _row("Cobertura", claim.cobertura)
    _row("Ramo", claim.ramo)
    _row("Ciudad / sucursal", f"{claim.ciudad} — {claim.sucursal}")
    _row("Fecha ocurrencia", str(claim.fecha_ocurrencia))
    _row("Fecha reporte", str(claim.fecha_reporte))
    _row("Monto reclamado", f"${claim.monto_reclamado:,.2f}")
    _row("Suma asegurada", f"${claim.suma_asegurada:,.2f}")
    _row("Estado del siniestro", claim.estado)
    if claim.proveedor:
        _row("Proveedor", claim.proveedor)
    if claim.vehiculo:
        v = claim.vehiculo
        _row(
            "Vehículo",
            f"{v.marca} {v.modelo} {v.anio} — placa {v.placa}",
        )

    doc.add_paragraph("")

    # -----------------------------------------------------------------------
    # 3. Score + nivel
    # -----------------------------------------------------------------------
    doc.add_heading("Score de riesgo", level=1)

    nivel_label = _NIVEL_LABELS.get(claim.nivel, claim.nivel.value)
    color = _NIVEL_COLORS.get(claim.nivel, (0, 0, 0))

    score_para = doc.add_paragraph()
    run_score = score_para.add_run(f"Score: {claim.score}/100  |  Nivel: {nivel_label}")
    run_score.bold = True
    run_score.font.color.rgb = RGBColor(*color)
    run_score.font.size = Pt(13)

    if claim.ml_probability is not None:
        doc.add_paragraph(
            f"Probabilidad ML de posible fraude: {claim.ml_probability * 100:.1f}%"
        )

    doc.add_paragraph("")

    # -----------------------------------------------------------------------
    # 4. Reglas activadas
    # -----------------------------------------------------------------------
    doc.add_heading("Reglas activadas", level=1)

    if not claim.alertas:
        doc.add_paragraph("Sin alertas activadas.")
    else:
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = "Código"
        hdr[1].text = "Puntos"
        hdr[2].text = "Detalle"
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.bold = True
        for alerta in claim.alertas:
            row_cells = table.add_row().cells
            row_cells[0].text = alerta.code
            row_cells[1].text = str(alerta.puntos)
            row_cells[2].text = alerta.detalle

    doc.add_paragraph("")

    # -----------------------------------------------------------------------
    # 4b. Análisis explicativo — ML / anomalía / NLP / panel multi-agente
    # -----------------------------------------------------------------------
    append_analysis_sections(doc, claim)

    # -----------------------------------------------------------------------
    # 5. Resumen
    # -----------------------------------------------------------------------
    doc.add_heading("Resumen del caso", level=1)

    resumen = claim.resumen_editado or _fallback_resumen(claim)
    doc.add_paragraph(resumen)

    doc.add_paragraph("")

    # -----------------------------------------------------------------------
    # 6. Recomendación
    # -----------------------------------------------------------------------
    doc.add_heading("Recomendación de revisión", level=1)
    recomendacion = _RECOMENDACION.get(
        claim.nivel,
        "Revisar el caso según los procedimientos establecidos.",
    )
    doc.add_paragraph(recomendacion)

    doc.add_paragraph("")
    doc.add_paragraph(
        "NOTA: Este sistema es una herramienta de apoyo. La decisión final "
        "corresponde siempre al analista humano de la Aseguradora del Sur.",
        style="Intense Quote",
    )

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _fallback_resumen(claim: ClaimDetail) -> str:
    """Data-driven fallback when no analyst-edited summary exists."""
    rules_fired = ", ".join(a.code for a in claim.alertas) or "ninguna"
    return (
        f"El siniestro {claim.id} fue reportado por {claim.asegurado} "
        f"el {claim.fecha_reporte}, con una fecha de ocurrencia del "
        f"{claim.fecha_ocurrencia}. El monto reclamado asciende a "
        f"${claim.monto_reclamado:,.2f} sobre una suma asegurada de "
        f"${claim.suma_asegurada:,.2f}. "
        f"El sistema asignó un score de {claim.score}/100 (nivel {claim.nivel.value}). "
        f"Reglas activadas: {rules_fired}. "
        "Este resumen es generado automáticamente; el analista puede editarlo."
    )


async def generate_claim_report_docx(claim: ClaimDetail) -> bytes:
    """Async entry point — offloads the blocking docx build to a thread."""
    return await asyncio.to_thread(_build_docx, claim)
