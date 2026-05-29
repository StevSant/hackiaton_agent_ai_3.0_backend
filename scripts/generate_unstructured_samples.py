"""Generate synthetic unstructured sample documents for the hackathon demo.

Produces 3 PDF and 3 Word (.docx) documents under data/casos_demo/.
Each document contains enough information for an LLM to extract a ClaimDetail.
All data is synthetic — NO real PII.

Run:
    uv run python scripts/generate_unstructured_samples.py

Outputs:
    data/casos_demo/pdf/denuncia_policial_robo_001.pdf      (SIN-DEMO-001, Toyota Hilux robo)
    data/casos_demo/pdf/informe_pericial_002.pdf            (SIN-DEMO-002, Chevrolet Aveo collision)
    data/casos_demo/pdf/boleta_siniestro_003.pdf            (SIN-DEMO-003, Kia Picanto docs incompletos)
    data/casos_demo/docx/informe_perito_004.docx            (SIN-DEMO-004, Hyundai Tucson rayon)
    data/casos_demo/docx/denuncia_005.docx                  (SIN-DEMO-005, Mazda CX-5 danos totales)
    data/casos_demo/docx/denuncia_robo_006.docx             (SIN-DEMO-006, Toyota Hilux robo similar)
"""

from __future__ import annotations

import sys
from pathlib import Path

# Progress output uses non-ASCII — force UTF-8 so it doesn't crash on Windows cp1252.
sys.stdout.reconfigure(encoding="utf-8")

# ---------------------------------------------------------------------------
# Report file constants
# ---------------------------------------------------------------------------
BASE = Path(__file__).parent.parent / "data" / "casos_demo"
# These unstructured fixtures are all vehiculos cases (classified by ramo).
PDF_DIR = BASE / "pdf" / "vehiculos"
DOCX_DIR = BASE / "docx" / "vehiculos"
PDF_DIR.mkdir(parents=True, exist_ok=True)
DOCX_DIR.mkdir(parents=True, exist_ok=True)


# ---------------------------------------------------------------------------
# PDF generation helpers (reportlab)
# ---------------------------------------------------------------------------

def _make_pdf(path: Path, title: str, sections: list[tuple[str, str]]) -> None:
    """Render a simple two-column-header document as PDF using reportlab."""
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "DocTitle",
        parent=styles["Heading1"],
        alignment=TA_CENTER,
        fontSize=14,
        spaceAfter=6,
    )
    subtitle_style = ParagraphStyle(
        "DocSubtitle",
        parent=styles["Normal"],
        alignment=TA_CENTER,
        fontSize=10,
        spaceAfter=12,
        textColor=(0.4, 0.4, 0.4),
    )
    section_title_style = ParagraphStyle(
        "SectionTitle",
        parent=styles["Heading2"],
        fontSize=11,
        spaceBefore=10,
        spaceAfter=4,
        textColor=(0.1, 0.1, 0.6),
    )
    body_style = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=9,
        leading=14,
        alignment=TA_LEFT,
    )

    story = [
        Paragraph("REPÚBLICA DEL ECUADOR", subtitle_style),
        Paragraph("ASEGURADORA DEL SUR S.A.", subtitle_style),
        Spacer(1, 0.3 * cm),
        HRFlowable(width="100%", thickness=1, color=(0.1, 0.1, 0.6)),
        Spacer(1, 0.3 * cm),
        Paragraph(title, title_style),
        Spacer(1, 0.5 * cm),
    ]

    for section_name, content in sections:
        story.append(Paragraph(section_name, section_title_style))
        story.append(Paragraph(content.replace("\n", "<br/>"), body_style))
        story.append(Spacer(1, 0.3 * cm))

    story.append(Spacer(1, 1 * cm))
    story.append(HRFlowable(width="100%", thickness=0.5, color=(0.6, 0.6, 0.6)))
    story.append(Paragraph(
        "Documento generado para fines de demostración — datos 100% sintéticos.",
        ParagraphStyle("Footer", parent=styles["Normal"], fontSize=7,
                       alignment=TA_CENTER, textColor=(0.5, 0.5, 0.5)),
    ))

    doc.build(story)
    print(f"  [PDF] {path}")


# ---------------------------------------------------------------------------
# Word document helpers (python-docx)
# ---------------------------------------------------------------------------

def _make_docx(path: Path, title: str, sections: list[tuple[str, str]]) -> None:
    """Render a simple document as .docx using python-docx."""
    import docx
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = docx.Document()

    # Narrow margins
    for section in document.sections:
        section.top_margin = docx.shared.Cm(2)
        section.bottom_margin = docx.shared.Cm(2)
        section.left_margin = docx.shared.Cm(2.5)
        section.right_margin = docx.shared.Cm(2.5)

    # Header block
    header_p = document.add_paragraph()
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header_p.add_run("REPÚBLICA DEL ECUADOR — ASEGURADORA DEL SUR S.A.")
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    title_p = document.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run(title)
    run.bold = True
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x99)

    document.add_paragraph()  # spacer

    for section_name, content in sections:
        heading = document.add_heading(section_name, level=2)
        heading.runs[0].font.color.rgb = RGBColor(0x1a, 0x1a, 0x99)
        body = document.add_paragraph(content)
        body.paragraph_format.space_after = Pt(6)

    document.add_paragraph()
    footer_p = document.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_p.add_run(
        "Documento generado para fines de demostración — datos 100% sintéticos."
    )
    run.font.size = Pt(7)
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

    document.save(str(path))
    print(f"  [DOCX] {path}")


# ---------------------------------------------------------------------------
# Document content definitions (one per caso)
# ---------------------------------------------------------------------------

def _generate_pdf_001() -> None:
    """caso_01 — Toyota Hilux robo total, Quito."""
    _make_pdf(
        PDF_DIR / "denuncia_policial_robo_001.pdf",
        "DENUNCIA FISCAL — ROBO DE VEHÍCULO",
        [
            ("DATOS DEL DENUNCIANTE", (
                "Nombre completo: Andrés Felipe Cordero Naranjo\n"
                "Cédula de identidad: 1703456789\n"
                "Código de asegurado: ASG-D0001\n"
                "Dirección: Calle Francisco de Orellana, sector Cumbayá, Quito\n"
                "Teléfono: 099-555-0201"
            )),
            ("DATOS DE LA PÓLIZA", (
                "Número de póliza: PV-DEMO-0001\n"
                "Ramo: Vehículos\n"
                "Cobertura contratada: Pérdida Total por Robo\n"
                "Suma asegurada: USD 42,000.00\n"
                "Vigencia: 12 de abril de 2026 al 12 de abril de 2027"
            )),
            ("DATOS DEL VEHÍCULO", (
                "Marca: Toyota\n"
                "Modelo: Hilux\n"
                "Año: 2023\n"
                "Placa: PCH-4821\n"
                "Número de chasis: MR0HZ8CD400123456\n"
                "Color: Gris oscuro"
            )),
            ("DESCRIPCIÓN DEL HECHO", (
                "El viernes 1 de mayo de 2026, aproximadamente a las 21:30 horas, el denunciante "
                "Andrés Felipe Cordero Naranjo reporta el robo de su Toyota Hilux 2023, "
                "placas PCH-4821, ocurrido en el sector de Cumbayá, calle Francisco de Orellana, "
                "mientras el vehículo se encontraba estacionado frente al restaurante "
                "La Tablita del Tártaro. Tres sujetos en motocicleta abordaron al asegurado "
                "al regresar a su vehículo, le exigieron las llaves bajo amenaza con arma de "
                "fuego y se dieron a la fuga en dirección norte con el vehículo. El asegurado "
                "quedó ileso y se trasladó por sus propios medios hasta la UPC del sector "
                "donde sentó el parte policial. La denuncia formal ante la Fiscalía General "
                "del Estado fue presentada el 8 de mayo de 2026 (siete días después del evento), "
                "dado que el denunciante se encontraba fuera del país por motivos laborales.\n"
                "Solicita activar la cobertura de Pérdida Total por Robo. "
                "Monto reclamado: USD 41,160.00."
            )),
            ("FECHA DE OCURRENCIA Y REPORTE", (
                "Fecha de ocurrencia: 01 de mayo de 2026\n"
                "Fecha de reporte a la aseguradora: 08 de mayo de 2026\n"
                "Número de caso Fiscalía: FGE-PICH-2026-004821"
            )),
            ("FIRMA DEL DENUNCIANTE", (
                "Andrés Felipe Cordero Naranjo\n"
                "C.I.: 1703456789\n"
                "[Firma manuscrita simulada]\n"
                "Quito, 08 de mayo de 2026"
            )),
        ],
    )


def _generate_pdf_002() -> None:
    """caso_02 — Chevrolet Aveo colisión, Guayaquil."""
    _make_pdf(
        PDF_DIR / "informe_pericial_002.pdf",
        "INFORME PERICIAL TÉCNICO — DAÑOS MATERIALES",
        [
            ("ENCABEZADO DEL PERITO", (
                "Perito asignado: Ing. Roberto Zamora Villafuerte\n"
                "Matrícula profesional: SPN-PER-00892\n"
                "Fecha de inspección: 13 de mayo de 2026\n"
                "Lugar de inspección: Taller Chapa y Pintura del Pacífico, Guayaquil"
            )),
            ("DATOS DEL SINIESTRO", (
                "Número de siniestro: SIN-DEMO-002\n"
                "Asegurado: Patricia Lorena Vásquez Montoya\n"
                "Cédula: 0912345678\n"
                "Código de asegurado: ASG-D0002\n"
                "Póliza: PV-DEMO-0002\n"
                "Cobertura: Daños Materiales\n"
                "Suma asegurada: USD 14,000.00\n"
                "Ciudad: Guayaquil"
            )),
            ("DATOS DEL VEHÍCULO", (
                "Marca: Chevrolet\n"
                "Modelo: Aveo\n"
                "Año: 2020\n"
                "Placa: GBR-2294\n"
                "Chasis: 8LBTZS5E9LG041872\n"
                "Kilometraje al momento del siniestro: 68,420 km"
            )),
            ("DESCRIPCIÓN DEL ACCIDENTE", (
                "El lunes 11 de mayo de 2026, aproximadamente a las 19:15 horas, en la Avenida "
                "Francisco de Orellana, intersección con Avenida Luis Orrantia, Guayaquil, "
                "el vehículo de la asegurada Patricia Lorena Vásquez Montoya fue embestido "
                "por detrás por un vehículo de color gris no identificado que se dio a la fuga "
                "tras el impacto. La asegurada se desplazaba con dirección al norte en tráfico "
                "moderado. La póliza de Daños Materiales fue contratada el 10 de mayo de 2026, "
                "un día antes de la ocurrencia del siniestro."
            )),
            ("DICTAMEN DE DAÑOS", (
                "Daños observados:\n"
                "- Parachoque posterior: deformación total, requiere reemplazo.\n"
                "- Tapa de maletero: abolladuras y deformación estructural, requiere reemplazo.\n"
                "- Luz trasera derecha: rotura completa, requiere reemplazo.\n"
                "- Carrocería trasera: abolladuras menores, requiere latonería y pintura.\n\n"
                "Monto estimado de reparación: USD 8,400.00\n"
                "Monto reclamado por la asegurada: USD 8,400.00\n"
                "Fecha de ocurrencia: 11 de mayo de 2026\n"
                "Fecha de reporte a la aseguradora: 12 de mayo de 2026"
            )),
            ("CONCLUSIÓN DEL PERITO", (
                "Los daños son consistentes con una colisión trasera a velocidad moderada. "
                "Se recomienda proceder con la liquidación conforme a la cobertura de Daños "
                "Materiales de la póliza PV-DEMO-0002.\n\n"
                "Ing. Roberto Zamora Villafuerte\n"
                "Perito Técnico Vehicular\n"
                "Guayaquil, 13 de mayo de 2026"
            )),
        ],
    )


def _generate_pdf_003() -> None:
    """caso_03 — Kia Picanto, Cuenca, docs incompletos."""
    _make_pdf(
        PDF_DIR / "boleta_siniestro_003.pdf",
        "BOLETA INTERNA DE SINIESTRO",
        [
            ("DATOS GENERALES", (
                "Número de siniestro: SIN-DEMO-003\n"
                "Fecha de apertura: 25 de abril de 2026\n"
                "Sucursal: Cuenca Centro\n"
                "Analista asignado: Lcda. Mónica Herrera"
            )),
            ("DATOS DEL ASEGURADO", (
                "Nombre: Rodrigo Sebastián Arias Cifuentes\n"
                "Cédula: 0103456789\n"
                "Código de asegurado: ASG-D0003\n"
                "Ciudad: Cuenca\n"
                "Teléfono: 072-555-0303"
            )),
            ("DATOS DE LA PÓLIZA", (
                "Número de póliza: PV-DEMO-0003\n"
                "Ramo: Vehículos\n"
                "Cobertura: Daños\n"
                "Suma asegurada: USD 8,000.00\n"
                "Vigencia: 01 de octubre de 2025 al 01 de octubre de 2026"
            )),
            ("VEHÍCULO SINIESTRADO", (
                "Marca: Kia\n"
                "Modelo: Picanto\n"
                "Año: 2019\n"
                "Placa: AZB-0571\n"
                "Chasis: KNABE211AK5306841"
            )),
            ("DESCRIPCIÓN DEL SINIESTRO", (
                "El asegurado reporta daños materiales en su Kia Picanto 2019, placas AZB-0571, "
                "ocurridos el 15 de abril de 2026 a las 11:00 horas en el estacionamiento del "
                "Centro Comercial Mall del Río, Cuenca. El vehículo presentó impacto lateral en "
                "la puerta delantera del conductor y guardabarro izquierdo, aparentemente causado "
                "por un vehículo que abandonó el lugar sin dejar datos. El asegurado demoró 10 "
                "días en reportar el siniestro debido a una emergencia familiar.\n"
                "Fecha de ocurrencia: 15 de abril de 2026\n"
                "Fecha de reporte: 25 de abril de 2026\n"
                "Monto reclamado: USD 3,200.00"
            )),
            ("ESTADO DE DOCUMENTACIÓN", (
                "Cédula de identidad: ENTREGADA\n"
                "Matrícula vehicular: ENTREGADA\n"
                "Parte policial / CTG: ENTREGADO\n"
                "Certificado médico de pasajeros: PENDIENTE\n"
                "Avalúo independiente del vehículo: PENDIENTE\n\n"
                "OBSERVACIÓN: El asegurado indicó que los documentos pendientes están en "
                "trámite y serán entregados en los próximos días hábiles."
            )),
            ("FIRMA DEL RECEPTOR", (
                "Lcda. Mónica Herrera\n"
                "Recepción de Siniestros — Sucursal Cuenca\n"
                "Cuenca, 25 de abril de 2026"
            )),
        ],
    )


def _generate_docx_004() -> None:
    """caso_04 — Hyundai Tucson raspón, Quito."""
    _make_docx(
        DOCX_DIR / "informe_perito_004.docx",
        "INFORME DE PERITO — DAÑOS PARCIALES",
        [
            ("DATOS DEL CASO", (
                "Número de siniestro: SIN-DEMO-004\n"
                "Perito: Ing. Carlos Albán Torres\n"
                "Fecha de inspección: 23 de abril de 2026\n"
                "Ciudad: Quito"
            )),
            ("DATOS DEL ASEGURADO", (
                "Nombre: Valeria Cristina Moreno Benítez\n"
                "Cédula: 1704567890\n"
                "Código de asegurado: ASG-D0004\n"
                "Póliza: PV-DEMO-0004\n"
                "Cobertura: Daños Parciales\n"
                "Suma asegurada: USD 28,000.00\n"
                "Vigencia: 01 de diciembre de 2025 al 01 de diciembre de 2026"
            )),
            ("DATOS DEL VEHÍCULO", (
                "Marca: Hyundai\n"
                "Modelo: Tucson\n"
                "Año: 2022\n"
                "Placa: PBD-1083\n"
                "Chasis: 5NMS24AJ6NH412095"
            )),
            ("DESCRIPCIÓN DEL SINIESTRO", (
                "La asegurada Valeria Cristina Moreno Benítez reporta un raspón superficial "
                "en el costado derecho de su Hyundai Tucson 2022, placas PBD-1083, producido "
                "en el estacionamiento del centro comercial Quicentro Shopping, Quito, "
                "el 20 de abril de 2026. Un tercero identificado dejó sus datos de contacto "
                "y reconoció la responsabilidad. El reclamo cubre únicamente trabajos de "
                "chapa y pintura en la puerta trasera derecha y el guardabarro.\n"
                "Fecha de ocurrencia: 20 de abril de 2026\n"
                "Fecha de reporte: 22 de abril de 2026\n"
                "Monto reclamado: USD 1,200.00"
            )),
            ("DICTAMEN", (
                "Daños observados:\n"
                "- Puerta trasera derecha: rayaduras superficiales, requiere pulido y pintura.\n"
                "- Guardabarro trasero derecho: rayaduras menores.\n\n"
                "Monto estimado de reparación: USD 1,200.00\n"
                "Documentación completa entregada. Caso de bajo riesgo.\n\n"
                "Ing. Carlos Albán Torres — Perito Vehicular Quito, 23 de abril de 2026"
            )),
        ],
    )


def _generate_docx_005() -> None:
    """caso_05 — Mazda CX-5 daños totales, Guayaquil."""
    _make_docx(
        DOCX_DIR / "denuncia_005.docx",
        "DENUNCIA DE SINIESTRO — DAÑOS TOTALES",
        [
            ("DATOS DEL DENUNCIANTE", (
                "Nombre: Jorge Mauricio Landívar Peñafiel\n"
                "Cédula: 0912378901\n"
                "Código de asegurado: ASG-D0005\n"
                "Ciudad: Guayaquil\n"
                "Teléfono: 098-555-0506"
            )),
            ("DATOS DE LA PÓLIZA", (
                "Número de póliza: PV-DEMO-0005\n"
                "Ramo: Vehículos\n"
                "Cobertura: Daños Totales\n"
                "Suma asegurada: USD 23,000.00\n"
                "Vigencia: 25 de abril de 2026 al 25 de abril de 2027"
            )),
            ("VEHÍCULO", (
                "Marca: Mazda\n"
                "Modelo: CX-5\n"
                "Año: 2021\n"
                "Placa: GBT-9917\n"
                "Chasis: JM3KFBCM5M0408316"
            )),
            ("DESCRIPCIÓN DEL SINIESTRO", (
                "El asegurado reporta daños severos en el lateral izquierdo de su Mazda CX-5 "
                "2021, placas GBT-9917, producidos en la madrugada del 2 de mayo de 2026 en "
                "un sector industrial al norte de Guayaquil, sin testigos presentes. "
                "El asegurado indica que encontró el vehículo con los daños al salir de una "
                "reunión de trabajo en las instalaciones de la empresa.\n"
                "Fecha de ocurrencia: 02 de mayo de 2026\n"
                "Fecha de reporte a la aseguradora: 12 de mayo de 2026\n"
                "Monto reclamado: USD 22,000.00"
            )),
            ("ESTADO DOCUMENTACIÓN", (
                "Cédula de identidad: ENTREGADA\n"
                "Matrícula vehicular: ENTREGADA\n"
                "Parte policial: ENTREGADO\n"
                "Avalúo independiente: PENDIENTE\n"
                "Fotos del daño: PENDIENTE\n\n"
                "NOTA: El asegurado fue informado de los documentos pendientes."
            )),
            ("FIRMA", (
                "Jorge Mauricio Landívar Peñafiel\n"
                "C.I.: 0912378901\n"
                "[Firma manuscrita simulada]\n"
                "Guayaquil, 12 de mayo de 2026"
            )),
        ],
    )


def _generate_docx_006() -> None:
    """caso_06 — Toyota Hilux robo similar a caso_01, Quito."""
    _make_docx(
        DOCX_DIR / "denuncia_robo_006.docx",
        "DENUNCIA FISCAL — ROBO DE VEHÍCULO",
        [
            ("DATOS DEL DENUNCIANTE", (
                "Nombre: Carlos Vinicio Espinoza Maldonado\n"
                "Cédula: 1705678901\n"
                "Código de asegurado: ASG-D0006\n"
                "Ciudad: Quito\n"
                "Teléfono: 099-555-0607"
            )),
            ("DATOS DE LA PÓLIZA", (
                "Número de póliza: PV-DEMO-0006\n"
                "Ramo: Vehículos\n"
                "Cobertura: Pérdida Total por Robo\n"
                "Suma asegurada: USD 40,000.00\n"
                "Vigencia: 15 de noviembre de 2025 al 15 de noviembre de 2026"
            )),
            ("VEHÍCULO ROBADO", (
                "Marca: Toyota\n"
                "Modelo: Hilux\n"
                "Año: 2022\n"
                "Placa: GBC-7142\n"
                "Chasis: MR0HZ8CD0N0387251\n"
                "Color: Negro"
            )),
            ("DESCRIPCIÓN DEL HECHO", (
                "El martes 12 de mayo de 2026, aproximadamente a las 23:00 horas, el denunciante "
                "Carlos Vinicio Espinoza Maldonado reporta el robo de su Toyota Hilux 2022, "
                "placas GBC-7142, ocurrido en el sector de La Floresta, calle Galavis, Quito, "
                "mientras el vehículo se encontraba estacionado frente al restaurante "
                "El Español de La Floresta. Dos sujetos en motocicleta interceptaron al asegurado "
                "al acercarse a su vehículo, le apuntaron con un arma de fuego y le exigieron "
                "las llaves, alejándose en dirección al parque La Carolina. El asegurado "
                "permaneció ileso y procedió a movilizarse a pie hasta la UPC más cercana.\n"
                "La denuncia formal ante la Fiscalía fue presentada el 19 de mayo de 2026, "
                "siete días después de la ocurrencia, dado que el asegurado se hallaba en "
                "comisión de servicios en la ciudad de Manta durante esa semana.\n"
                "Monto reclamado: USD 38,500.00"
            )),
            ("FECHAS", (
                "Fecha de ocurrencia: 12 de mayo de 2026\n"
                "Fecha de reporte a la aseguradora: 19 de mayo de 2026\n"
                "Número de caso Fiscalía: FGE-PICH-2026-007142"
            )),
            ("FIRMA", (
                "Carlos Vinicio Espinoza Maldonado\n"
                "C.I.: 1705678901\n"
                "[Firma manuscrita simulada]\n"
                "Quito, 19 de mayo de 2026"
            )),
        ],
    )


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    print("Generando documentos de muestra sintéticos...")
    _generate_pdf_001()
    _generate_pdf_002()
    _generate_pdf_003()
    _generate_docx_004()
    _generate_docx_005()
    _generate_docx_006()
    print("Listo. Documentos generados en data/casos_demo/pdf/ y data/casos_demo/docx/")


if __name__ == "__main__":
    main()
